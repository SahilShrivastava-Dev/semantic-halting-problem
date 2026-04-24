"""
generate_docs.py

Auto-generates the SHP_Codebase_Guide.docx documentation for the
Semantic Halting Problem (SHP) project.

Run from the project root:
    python doc/generate_docs.py

The output file is written to doc/SHP_Codebase_Guide.docx.

Dependencies:
    pip install python-docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ─────────────────────────────────────────────────────────────
# Formatting helpers
# ─────────────────────────────────────────────────────────────
def _heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading at the given outline level."""
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    palette = {1: (0x1A, 0x73, 0xE8), 2: (0x18, 0x45, 0x9B), 3: (0x2E, 0x7D, 0x32)}
    r, g, b = palette.get(level, (0x33, 0x33, 0x33))
    run.font.color.rgb = RGBColor(r, g, b)
    return h


def _code_block(doc: Document, code: str):
    """Add a shaded monospace code block paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return para


def _info_box(doc: Document, label: str, text: str):
    """Add a 💡/ℹ️ callout paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    r_label = para.add_run(f"{label}  ")
    r_label.bold = True
    r_label.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    r_text = para.add_run(text)
    r_text.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _table(doc: Document, style: str, headers: list[str], rows: list[tuple]) -> None:
    """Add a formatted two-column table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val


# ─────────────────────────────────────────────────────────────
# Document builder
# ─────────────────────────────────────────────────────────────
def build_doc() -> None:
    """Construct and save the complete codebase guide document."""
    doc = Document()

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Semantic Halting Problem (SHP)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    sub = doc.add_paragraph("Complete Production Codebase Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph("")
    intro = doc.add_paragraph(
        "This document walks through every file, class, and function in the SHP project "
        "in plain English.  No prior AI or machine-learning experience is required.  "
        "By the end you will understand what the project does, why each file exists, "
        "and how all the pieces connect."
    )
    intro.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ── Section 1: What is SHP? ──────────────────────────────────────────────
    _heading(doc, "1.  What is the Semantic Halting Problem?", level=1)
    doc.add_paragraph(
        "Imagine two AI agents — a Writer and a Critic — collaborating on a report.  "
        "The Writer produces a draft, the Critic reviews it and gives feedback, and the "
        "Writer revises.  In theory they finish when the report is perfect.  In practice "
        "a real LLM Critic eventually runs out of substantive critique and starts "
        "recycling minor wording suggestions — the loop is stuck in a semantic deadlock."
    )
    doc.add_paragraph(
        "The Semantic Halting Problem (SHP) asks: how can the computer detect and halt "
        "this deadlock automatically, without a hard-coded step limit?  "
        "This project's answer is mathematical: convert each draft to a vector "
        "(embedding), measure the cosine distance between consecutive drafts, and stop "
        "when the distance collapses to near-zero — meaning no new information is being "
        "generated."
    )
    _info_box(doc, "💡 Key Idea:",
              "Distance ≈ 0 means 'no new information'. That is the halt signal.")

    doc.add_paragraph("")
    _heading(doc, "2.  Project Architecture at a Glance", level=1)
    doc.add_paragraph(
        "The project consists of seven Python files, each with a single, clear "
        "responsibility:"
    )
    _table(doc, "Light List Accent 1",
           ["File", "Responsibility"],
           [
               ("config.py",               "Central constants, file paths, and model identifiers."),
               ("main.py",                 "Five-phase pipeline orchestrator."),
               ("agent_workflow.py",       "LangGraph multi-agent graph (Writer→Evaluator→Critic→halt)."),
               ("agents.py",               "Real LLM-powered Writer and Critic nodes (Groq/llama-3.1-8b)."),
               ("semantic_entropy.py",     "Cosine-distance convergence engine."),
               ("ragas_eval.py",           "Post-hoc Ragas evaluation of final agent outputs."),
               ("optimize_score.py",       "Linear Regression to learn IS metric weights."),
               ("test_information_score.py","IS formula validation (good > convergent > poor)."),
               ("doc/generate_docs.py",    "This documentation generator."),
           ])
    doc.add_page_break()

    # ── Section 3: File-by-File Breakdown ────────────────────────────────────
    _heading(doc, "3.  File-by-File Breakdown", level=1)

    # config.py
    _heading(doc, "3.1  config.py  (Central Configuration)", level=2)
    doc.add_paragraph(
        "All tunable constants live here: model names, halting thresholds, file paths, "
        "and default weights.  Every other module imports from config.py, so changing "
        "a constant once propagates everywhere."
    )
    _code_block(doc,
        "# Key constants:\n"
        "CONVERGENCE_THRESHOLD = 0.06   # cosine distance below which we halt\n"
        "MAX_ROUNDS            = 12     # hard failsafe cap\n"
        "AGENT_LLM_MODEL       = 'llama-3.1-8b-instant'  # Groq model\n"
        "EMBEDDING_MODEL_NAME  = 'BAAI/bge-small-en-v1.5'"
    )

    # semantic_entropy.py
    _heading(doc, "3.2  semantic_entropy.py  (Math Engine)", level=2)
    doc.add_paragraph(
        "The mathematical backbone of SHP.  Wraps any Langchain-compatible embedding "
        "model and exposes two methods:"
    )
    _heading(doc, "get_embedding(text) → List[float]", level=3)
    doc.add_paragraph(
        "Converts a string to a dense 384-dimensional vector "
        "(BAAI/bge-small-en-v1.5).  Similar texts produce vectors pointing "
        "in nearly the same direction."
    )
    _heading(doc, "calculate_distance(v1, v2) → float", level=3)
    doc.add_paragraph(
        "Computes cosine distance (1 − cosine_similarity).  "
        "Returns 0.0 for identical texts, 1.0 for orthogonal, up to 2.0 for "
        "opposite.  Synonym rewrites score ~0.04; the halt threshold is 0.06."
    )
    _code_block(doc,
        "dist = calc.calculate_distance(vec_draft_4, vec_draft_5)\n"
        "# 0.003 → semantically stuck → HALT\n"
        "# 0.25  → meaningful revision → CONTINUE"
    )

    # agents.py
    _heading(doc, "3.3  agents.py  (LLM-Powered Agents)", level=2)
    doc.add_paragraph(
        "Defines the Writer and Critic as real LLM agent nodes powered by "
        "Groq's llama-3.1-8b-instant (free tier, fast inference).  Both share a "
        "single ChatGroq instance loaded once at module import."
    )
    _heading(doc, "writer_node(state) → dict", level=3)
    doc.add_paragraph(
        "On round 0: generates the initial draft from the scenario topic and brief.  "
        "On subsequent rounds: rewrites the draft to address the Critic's feedback, "
        "producing a genuinely improved revision rather than a cosmetic paraphrase."
    )
    _heading(doc, "critic_node(state) → dict", level=3)
    doc.add_paragraph(
        "Reviews the current draft and returns exactly ONE substantive critique.  "
        "Returns the word APPROVED when no meaningful gaps remain — this triggers "
        "the Critic Approval halt signal independently of the entropy check."
    )
    _info_box(doc, "🔑 Why real LLMs?",
              "Real models organically exhaust their critique vocabulary across rounds, "
              "causing semantic convergence to emerge naturally without scripting.")

    # agent_workflow.py
    _heading(doc, "3.4  agent_workflow.py  (LangGraph Orchestrator)", level=2)
    doc.add_paragraph(
        "Assembles the LangGraph StateGraph and executes the per-scenario loop.  "
        "Contains three node functions and one conditional edge function."
    )
    _heading(doc, "Graph topology", level=3)
    _code_block(doc,
        "writer → evaluator → embed_state → check_convergence ──► critic ─► (back to writer)\n"
        "                                                       └──► END"
    )
    _heading(doc, "evaluator_node  — real-time IS scoring", level=3)
    doc.add_paragraph(
        "Runs Ragas inside the loop after every draft.  Applies learned IS weights "
        "to produce a per-round Information Score so the gain can be tracked in real time."
    )
    _heading(doc, "check_convergence  — the halting oracle", level=3)
    doc.add_paragraph(
        "Evaluates four signals in priority order: "
        "(1) Critic approval, (2) semantic entropy, (3) IS-Gain ≤ 0, (4) failsafe cap.  "
        "Returns 'end' on the first triggered signal."
    )

    # ragas_eval.py
    _heading(doc, "3.5  ragas_eval.py  (Post-Hoc Grader)", level=2)
    doc.add_paragraph(
        "After agent_workflow.py saves agent_results.json, ragas_eval.py "
        "re-evaluates every final draft with the full Ragas suite and saves "
        "ragas_scores.json for optimize_score.py."
    )
    _heading(doc, "The Four Metrics", level=3)
    _table(doc, "Light List Accent 2",
           ["Metric", "Plain English Meaning"],
           [
               ("Faithfulness",      "Did the answer hallucinate? Every claim must be supported by contexts."),
               ("Answer Relevancy",  "Does the answer actually address the question asked?"),
               ("Context Precision", "Of the retrieved contexts, what fraction was useful?"),
               ("Context Recall",    "Did the contexts cover everything needed to answer correctly?"),
           ])

    # optimize_score.py
    _heading(doc, "3.6  optimize_score.py  (IS Weight Optimizer)", level=2)
    doc.add_paragraph(
        "Fits scikit-learn LinearRegression(fit_intercept=False) to learn the "
        "weights w₁…w₄ that make IS the best single predictor of overall quality.  "
        "Coefficients are clipped to ≥ 0 and normalised to sum to 1."
    )
    _code_block(doc,
        "IS = 0.39·Faithfulness + 0.30·AnswerRelevancy\n"
        "   + 0.09·ContextPrecision + 0.23·ContextRecall\n\n"
        "R² = 0.81  ← strong fit (explains 81% of quality variance)"
    )
    _info_box(doc, "💡 Note:",
              "In production, replace the equal-weight proxy with real human annotations "
              "to obtain domain-specific weights.")

    # test_information_score.py
    _heading(doc, "3.7  test_information_score.py  (IS Validation)", level=2)
    doc.add_paragraph(
        "Validates the IS formula by scoring three quality tiers for each scenario "
        "and asserting IS(good) > IS(convergent) > IS(poor).  This confirms the "
        "formula discriminates between meaningful progress and deadlock output."
    )

    doc.add_page_break()

    # ── Section 4: Pipeline ──────────────────────────────────────────────────
    _heading(doc, "4.  Running the Full Pipeline", level=1)
    steps = [
        ("Activate venv",              "source venv/bin/activate"),
        ("Install dependencies",       "pip install -r requirements.txt"),
        ("Run the full pipeline",      "python main.py"),
        ("Or run phases individually",
         "python agent_workflow.py --split train\n"
         "python ragas_eval.py\n"
         "python optimize_score.py\n"
         "python agent_workflow.py --split val\n"
         "python test_information_score.py"),
    ]
    for label, cmd in steps:
        doc.add_paragraph(label, style="List Number")
        _code_block(doc, cmd)

    doc.add_paragraph("")
    _heading(doc, "5.  Glossary", level=1)
    _table(doc, "Light List Accent 1",
           ["Term", "Meaning"],
           [
               ("Embedding",          "A list of numbers capturing the meaning of text so distances can be computed."),
               ("Cosine Distance",    "Angle-based similarity measure between two vectors (0=identical, 1=orthogonal)."),
               ("LangGraph",          "Python library for building AI agent workflows as directed graphs."),
               ("Ragas",              "Open-source framework for evaluating RAG (Retrieval-Augmented Generation) quality."),
               ("Information Score",  "Weighted composite of the four Ragas metrics; the primary halting signal."),
               ("Linear Regression",  "ML algorithm that learns a linear formula to predict an output from inputs."),
               ("R² Score",           "Goodness-of-fit measure (0–1). Higher = formula explains more quality variance."),
               ("Convergence",        "The point where drafts stop changing semantically — the halt trigger."),
               ("Groq",               "Cloud inference API providing free, fast LLaMA-3 inference without a local GPU."),
               ("venv",               "Isolated Python environment for this project's dependencies."),
           ])

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = "doc/SHP_Codebase_Guide.docx"
    doc.save(out_path)
    print(f"✅  Documentation saved → {out_path}")


if __name__ == "__main__":
    build_doc()
